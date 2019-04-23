from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.core.files.storage import FileSystemStorage
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from datetime import date

fs = FileSystemStorage(location='/media')                   # to set the default storage path

# these classes are file managers and the variables are attributes which are stored in database



def eng_to_nep():
    a = date.today()
    y = int(a.year)
    m = int(a.month)
    d = int(a.day)

    d = 16 + d

    if d > 30:
        m = m + 9
        d = d % 30
    else:
        m = m + 8

    if m > 12:
        y = y + 57
        m = m % 12
    else:
        y = y + 56

    return str(y) + '-' + str(m) + "-" + str(d)

# receiver tag brings the execution here no matter where it is


def create_doc(class_no, sem_name, teacher_name, dep_name, records):

    if 'Prof. Dr.' in teacher_name:
        post = 'Professor'
    else:
        post = 'Teacher'

    document = Document()

    #Different styles
    obj_charstyle = document.styles.add_style('nepali', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Preeti'

    obj_charstyle = document.styles.add_style('parnepali', WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Preeti'

    obj_charstyle = document.styles.add_style('english', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(11)
    obj_font.name = 'Calibri Body'

    obj_font = document.styles['Normal'].font
    obj_font.name='Calibri Body'

    #header at top
    paragraph=document.add_paragraph()
    paragraph.style='parnepali'
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run('lqe\'jg ljZjljBfno\n').bold=True
    paragraph.add_run('OlGhlgol/Gª cWoog ;+:yfg\n').bold=True
    paragraph.add_run('s]lGb|o SofDk; k\'Nrf]s\n').bold=True
    paragraph.add_run('lzIfs sfo{ ;Dkfbg kmf/fd').bold=True

    #Introduction table

    table1=['ljefu ','Semester ','lzIfssf] gfd y/ ','kb ','k|lt xKtf sIff ePsf] lbg ']
    table2=[dep_name, sem_name, teacher_name, post, str(class_no)]
    table=document.add_table(3,2,style = 'Table Grid')
    table.autofit=True
    table.cell(0,0).width=Inches(6)
    table.cell(0,1).width=Inches(2)
    table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for i in range(int(len(table1))):
        row=table.rows[int(i/2)]
        paragraph=row.cells[i%2].paragraphs[0]
        if i%2==1:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if i==1 :
            paragraph.add_run(table1[i]+" : ").bold=True
        else:
            paragraph.add_run(table1[i]+" M ",style='nepali').bold=True
        paragraph.add_run(table2[i]).bold=True
    paragraph=document.add_paragraph()
    paragraph.add_run("\n:gfts tx sIff ljj/0f",style='nepali').underline=True

    #our main table

    table1=['qm=;+=','ljifo','sIff lsl;d',';+nUg lzIfs ;+Vof','lkl/o8','ljBfyL{ ;+Vof']
    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.autofit =True

    table.cell(0,0).width=Inches(1)
    table.cell(0,1).width=Inches(14)
    table.cell(0,2).width=Inches(2)
    table.cell(0,3).width=Inches(5)
    table.cell(0,4).width=Inches(7)
    table.cell(0,5).width=Inches(1)

    for i in range(len(table1)):
        cell=table.cell(0,i)
        paragraph=cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(table1[i],style='nepali').bold=True

    total = 0

    for snid, subid, classid, teacherid, periodid, studentid in records:
        row_cells = table.add_row().cells
        for i in range(5):
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        s=row_cells[0].paragraphs[0].add_run(snid)
        s.bold=True
        s.font.size=Pt(10)
        s=row_cells[1].paragraphs[0].add_run(subid)
        s.bold=True
        s.font.size=Pt(10)
        s=row_cells[2].paragraphs[0].add_run(classid)
        s.bold=True
        s.font.size=Pt(10)
        s=row_cells[3].paragraphs[0].add_run(teacherid)
        s.bold=True
        s.font.size=Pt(10)
        s=row_cells[4].paragraphs[0].add_run(periodid)
        s.bold=True
        s.font.size=Pt(10)
        s=row_cells[5].paragraphs[0].add_run(studentid)
        s.bold=True
        s.font.size=Pt(10)
        total = float(total) + float(periodid)
    cell=table.add_row().cells
    paragraph="Total = "+str(total)+" Periods"
    cell[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    s=cell[4].paragraphs[0].add_run(paragraph)
    s.bold=True
    s.font.size=Pt(9)
    paragraph=document.add_paragraph()
    paragraph.add_run("\n:gftsf]Q/ tx sIff ljj/0f",style='nepali').underline=True

    #Master level Table
    table1=['qm=;+=','ljifo',';+nUg lzIfs ;+Vof','q]ml86 cfj/']
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.autofit = True
    table.cell(0,1).width=Inches(4)
    for i in range(len(table1)):
        cell=table.cell(0,i).paragraphs[0]
        cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.add_run(table1[i],style='nepali').bold=True
    table.add_row()

    #Notes and periods
    s = document.add_paragraph()
    s.add_run("\nGff]6 M    -s_ lkl/o8 M kf7|oqmddf pNn]v eP adf]lhd x'g]5 .\n",style='nepali')
    s.add_run(" \t-v_ sIff lsl;d eGgfn]",style='nepali')
    s.add_run(" 1,2,3 ").bold=True
    s.add_run("hgfpg' kg]{5 .\n",style='nepali')
    s.add_run("\t   1)").bold=True
    s.add_run("eGgfn]",style='nepali')
    s.add_run(" Theory /Tutorial /B.E. Project /B.Arch. Thesis\n\t   2)").bold=True
    s.add_run(" eGgfn]",style='nepali')
    s.add_run(" Drawing /Design /Design Studio /Paper work").bold=True
    s.add_run(" x'g] Nofj\n",style='nepali')
    s.add_run("\t   3)").bold=True
    s.add_run(" eGgfn] ",style='nepali')
    s.add_run("2").bold=True
    s.add_run(" df pNn]v gePsf Nofjx?",style='nepali')

    document.add_paragraph()
    #document.add_page_break()

    table1=['lzIfssf] x:tfIf/ M =========================================','k|dfl0ft ug]{','lzIfssf] gfd y/ M ','laefuLo k|d\'v','ldtL M ']
    #table2=["","", teacher.name, "", str(date.today())]
    table2=["","", teacher_name, "", eng_to_nep()]



    table=document.add_table(3,2)
    table.autofit=True
    table.cell(0,0).width=Inches(10)
    table.cell(0,1).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
    table.cell(1,1).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
    for i in range(len(table1)):
        cell=table.cell(int(i/2),i%2)
        paragraph=cell.paragraphs[0]
        paragraph.add_run(table1[i],style='nepali')
        paragraph.add_run(table2[i]).bold=True

    #footnote
    paragraph = document.add_paragraph()
    paragraph.page_break_before = True
    paragraph.style='parnepali'
    paragraph.add_run("\n\n\n\nb|i6Jo M\t!_ sIff ?l6g ;+nUg x'g''kg]{5 .	 @_ ")
    paragraph.add_run("Elective Course ",style='english').bold=True
    paragraph.add_run("sf nflu ljBfyL{ ;+Vof $* x'g]5 .\n\t#_")
    paragraph.add_run(" Master/Ph.D. ",style='english').bold=True
    paragraph.add_run("sf]")
    paragraph.add_run(" Thesis ",style='english').bold=True
    paragraph.add_run("sf nflu of] kmd{ eg{ cfjZos 5}g .")

    document.add_page_break()
    # saves teacher name in database in dsauser model
    # saves document with the teacher name as file name
    teachername = teacher_name.replace(' ', '')

    document.save(str(teachername.lower()) + '.docx')

    f = open(str(teachername.lower()) + '.docx', 'rb')


    return f






